"""
Generate Pathogen_Economy_Epiforecast_Year1_Operations_Handbook.docx
(12-month budget alignment + how each line maps to this codebase).
Run from project root: python scripts/build_year1_operations_handbook_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Pathogen_Economy_Epiforecast_Year1_Operations_Handbook.docx"

BUDGET_ROWS = [
    (
        "P1",
        "Domain Registration",
        "Cloudflare",
        "1",
        "39,000",
        "468,000",
        "Core domain ownership: PathogenEconomyEpiforecast.com (STI intellectual property).",
    ),
    (
        "P1",
        "DNS + WAF/CDN",
        "Cloudflare",
        "1",
        "74,000",
        "888,000",
        "Edge DNS, WAF, and CDN in front of the Industry 4.0+ origin server.",
    ),
    (
        "P1",
        "AI Provider (Primary)",
        "OpenAI",
        "1",
        "2,260,503",
        "27,126,034",
        "Powers forecasting, NLP alerts, and AI signal validation.",
    ),
    (
        "P2",
        "AI Failover",
        "Groq",
        "1",
        "541,666",
        "6,500,000",
        "Backup LLM when OpenAI is unavailable or rate-limited.",
    ),
    (
        "P1",
        "Transactional Email",
        "SendGrid",
        "1",
        "66,666",
        "800,000",
        "Password resets, risk bulletins, and operational mail with strong deliverability.",
    ),
    (
        "P1",
        "Uptime + Paging",
        "Better Stack",
        "1",
        "100,000",
        "1,200,000",
        "Synthetic monitoring, SMS/voice escalation to on-call IT staff.",
    ),
    (
        "P2",
        "Backup Storage",
        "Cloudflare R2",
        "1",
        "36,666",
        "440,000",
        "Off-site object storage for database and file backups.",
    ),
    (
        "P1",
        "App Builder License",
        "ABQ",
        "1",
        "0",
        "0",
        "Free-trial tier in year 1 for complementary UI / deployment workflows.",
    ),
    (
        "P1",
        "Technical Staff",
        "IT Personnel",
        "2",
        "550,000",
        "13,200,000",
        "Two stipended engineers for monitoring, patching, and incident response.",
    ),
]


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_budget_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=7)
    hdr = table.rows[0].cells
    headers = [
        "Priority",
        "Item",
        "Vendor",
        "Qty",
        "Monthly (UGX)",
        "Year 1 Total (UGX)",
        "Notes",
    ]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    for row in BUDGET_ROWS:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    total_row = table.add_row().cells
    total_row[0].text = "TOTAL"
    total_row[1].text = ""
    total_row[2].text = ""
    total_row[3].text = ""
    total_row[4].text = "3,730,836"
    total_row[5].text = "44,770,034"
    total_row[6].text = "Complete 12-month operational and development budget."
    for p in total_row[0].paragraphs:
        for r in p.runs:
            r.bold = True

    doc.add_paragraph()


def main() -> None:
    doc = Document()
    _add_title(doc, "Pathogen Economy Epiforecast")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Year 1 Operations Handbook — Budget Alignment & System Behaviour")
    r.italic = True
    r.font.size = Pt(12)
    doc.add_paragraph()

    _add_heading(doc, "1. Executive summary", level=1)
    doc.add_paragraph(
        "This handbook ties the approved twelve-month operating budget to concrete responsibilities "
        "in the Pathogen Economy Epiforecast platform (Streamlit national dashboard, FastAPI companion API, "
        "SQLite validated-signal store, and Expo mobile client). Each budget line below lists what the "
        "vendors provide, how the application consumes those services today, and what IT staff must operate."
    )

    _add_heading(doc, "2. Approved budget (UGX)", level=1)
    _add_budget_table(doc)

    _add_heading(doc, "3. How the application uses each budget line", level=1)

    sections = [
        (
            "P1 — Domain (Cloudflare Registrar)",
            "The production hostname (e.g. pathogeneconomyepiforecast.com) establishes STI-controlled "
            "intellectual property and stable URLs for the dashboard and API documentation. "
            "DNS records eventually point at the Industry 4.0+ origin or a reverse proxy tier.",
        ),
        (
            "P1 — DNS + WAF / CDN (Cloudflare)",
            "Cloudflare terminates TLS, performs bot mitigation, caches static assets, and shields the upstream "
            "Streamlit / FastAPI processes. The Python code does not embed Cloudflare APIs; configuration lives "
            "in Cloudflare and on the host. Rate limiting inside FastAPI (`API_RATE_LIMIT_PER_MIN`) adds a second "
            "layer of request protection.",
        ),
        (
            "P1 — AI primary (OpenAI)",
            "Configured through `AI_API_KEY`, `AI_BASE_URL`, and `AI_MODEL` (or `OPENAI_*` aliases). "
            "OpenAI drives: (a) batched disease-signal validation in `signal_validator.py`, (b) Forecast Lab "
            "JSON briefs in `forecast_lab_four_disease.py`, (c) optional NLP alert copy in `data_services.py`, "
            "and (d) OpenAI-compatible proxy routes in `api_server.py` (`/v1/chat/completions`, `/v1/nlp-alerts`, "
            "`/v1/cursor/chat`).",
        ),
        (
            "P2 — AI failover (Groq)",
            "Set `AI_FAILOVER_API_KEY` or `GROQ_API_KEY` plus optional base/model overrides. The codebase now "
            "walks the provider chain automatically: primary credentials first, then Groq. The same chain is used "
            "for direct HTTP validation batches, `chat_text_from_messages` in `ai_config.py`, NLP alert generation, "
            "and the FastAPI proxy (environment credentials only — per-request Bearer overrides stay on a single host).",
        ),
        (
            "P1 — Transactional email (SendGrid)",
            "`send_admin_email` in `data_services.py` sends daily and emergency epidemic bulletins configured under "
            "Administration → risk email routing (`app_pages.py`). If `SENDGRID_API_KEY` is present, SMTP defaults "
            "to `smtp.sendgrid.net` with username `apikey`. Generic SMTP remains supported for other providers.",
        ),
        (
            "P1 — Uptime + paging (Better Stack)",
            "Better Stack performs external synthetic checks and on-call notifications. Target the public `/health` "
            "endpoint on `api_server.py` (returns `pathogen-economy-epiforecast-api` metadata) plus the public "
            "Streamlit URL once exposed. Optional `BETTER_STACK_HEARTBEAT_URL` can be set for cron-style heartbeats "
            "from backup jobs.",
        ),
        (
            "P2 — Backup storage (Cloudflare R2)",
            "Critical artefacts include `SIGNAL_DB_PATH` (validated SQLite feed), `uploads/reports` (Admin uploads), "
            "and environment templates. R2 is S3-compatible: IT should schedule encrypted nightly `aws s3 sync` "
            "(or rclone) jobs using `R2_*` credentials referenced in `.env.example`. The Admin panel surfaces "
            "`R2_BUCKET_NAME` / access-key presence as a checklist item only — uploads are still operated by staff.",
        ),
        (
            "P1 — App Builder license (ABQ trial)",
            "ABQ or similar low-code tooling can wrap deployment UX while this repository remains the source of "
            "truth for analytics logic. Track trial conversion before year two. Set `ABQ_PROJECT_ID` when the "
            "integration is live so administrators see readiness in-app.",
        ),
        (
            "P1 — Technical staff (two IT experts)",
            "Operate vendor consoles (Cloudflare, OpenAI, Groq, SendGrid, Better Stack, R2), rotate API keys, watch "
            "cost dashboards, restore backups, patch dependencies (`requirements.txt`), and coordinate with Industry 4.0+ "
            "for origin capacity. They should maintain runbooks for: failover drills (disable primary LLM key), email "
            "bounce handling, and incident communications to STI leadership.",
        ),
    ]

    for title, body in sections:
        _add_heading(doc, title, level=2)
        doc.add_paragraph(body)

    _add_heading(doc, "4. Component map (repository)", level=1)
    doc.add_paragraph(
        "• app.py / app_pages.py / pathogen_economy_pages.py — Streamlit UX, navigation, admin controls.\n"
        "• api_server.py — FastAPI NLP alerts, SEIR endpoint, OpenAI-compatible proxy, `/health` for monitors.\n"
        "• data_services.py — Data ingestion, SMTP alerts, epidemiology helpers.\n"
        "• ai_config.py — Provider resolution + automatic Groq failover for chat-style calls.\n"
        "• signal_validator.py — LLM batches with failover for outbreak signal adjudication.\n"
        "• signal_store.py — SQLite persistence for validated signals.\n"
        "• mobile-app/ — Expo client aligned to the same API patterns.\n"
        "• .env.example — authoritative list of integration environment variables."
    )

    _add_heading(doc, "5. Acceptance criteria for go-live", level=1)
    doc.add_paragraph(
        "1. Primary OpenAI credentials verified via Forecast Lab and a signal-heavy dashboard refresh.\n"
        "2. Groq failover verified by revoking the primary key in a staging environment — requests should still succeed.\n"
        "3. SendGrid (or SMTP) test bulletin from Administration succeeds for all leadership recipients.\n"
        "4. Better Stack monitors green for `/health` and the public dashboard URL for at least 72 hours.\n"
        "5. R2 backup job restored successfully into a scratch VM (tabletop exercise).\n"
        "6. Two named IT operators documented in the on-call roster with access to all vendor consoles."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
