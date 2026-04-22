"""
Build STI-EPI-FORECAST_Illustrated_Government_Handbook.docx with embedded figures.

Installs doc build deps on first run (python-docx, matplotlib) then generates
clean concept diagrams — not live browser screenshots, but professional visuals
for government briefings.
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path
from tempfile import TemporaryDirectory


def _ensure_deps() -> None:
    try:
        import docx  # noqa: F401
        import matplotlib  # noqa: F401
    except ImportError:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "-q", "python-docx", "matplotlib"],
        )


def _fig_architecture(out: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    boxes = [
        (1, 4.2, 2.2, 1.1, "Streamlit\nDashboard\n(app.py)"),
        (4, 4.2, 2.2, 1.1, "FastAPI\n(api_server.py)"),
        (7, 4.2, 2.2, 1.1, "Expo\nMobile"),
        (1, 2.2, 2.2, 1.1, "Data layer\n(data_services.py)"),
        (4, 2.2, 2.2, 1.1, "Open-web\nfeeds"),
        (7, 2.2, 2.2, 1.1, "AI / SMTP\nalerts"),
        (2.5, 0.5, 5, 1, "Ministry of Health Uganda — decision support"),
    ]
    for x, y, w, h, t in boxes:
        ax.add_patch(
            plt.Rectangle(
                (x, y),
                w,
                h,
                fill=True,
                facecolor="#0f172a",
                edgecolor="#22c55e",
                linewidth=2,
            )
        )
        ax.text(x + w / 2, y + h / 2, t, ha="center", va="center", color="white", fontsize=10, weight="bold")
    ax.set_title("STI-EPI-FORECAST — system architecture (concept)", color="#0f172a", fontsize=14, weight="bold", pad=12)
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="#f1f5f9")
    plt.close(fig)


def _fig_data_flow(out: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 4))
    ax.axis("off")
    sources = ["GDELT", "Reddit", "Hacker News", "NewsAPI\n(optional)"]
    xs = [1, 3, 5, 7]
    for x, s in zip(xs, sources):
        ax.add_patch(plt.Circle((x, 2), 0.55, color="#1e3a5f", ec="#38bdf8", lw=2))
        ax.text(x, 2, s, ha="center", va="center", color="white", fontsize=9, weight="bold")
    ax.arrow(7.8, 2, 1.2, 0, head_width=0.15, head_length=0.15, fc="#22c55e", ec="#22c55e", lw=2)
    ax.add_patch(plt.Rectangle((9.2, 1.35), 2.6, 1.3, fill=True, facecolor="#14532d", ec="#22c55e", lw=2))
    ax.text(10.5, 2, "Outbreak\nsnapshot", ha="center", va="center", color="white", fontsize=10, weight="bold")
    ax.set_xlim(0, 12.5)
    ax.set_ylim(0.5, 3.5)
    ax.set_title("Open-web signals → unified outbreak snapshot", fontsize=13, weight="bold", pad=8)
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="#f8fafc")
    plt.close(fig)


def _fig_roles(out: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(9, 5))
    ax.axis("off")
    roles = [
        "Incident\nCommander",
        "Surveillance\nAnalyst",
        "Epidemiology\nModeler",
        "Border\nOperations",
        "Policy &\nInvestment",
        "System\nAdministrator",
    ]
    angles = [i * 60 for i in range(6)]
    import math

    cx, cy, r = 4.5, 2.8, 2.0
    for i, (ang, label) in enumerate(zip(angles, roles)):
        rad = math.radians(ang - 90)
        x, y = cx + r * math.cos(rad), cy + r * math.sin(rad)
        ax.add_patch(plt.Circle((x, y), 0.62, color="#1e293b", ec="#f59e0b", lw=2))
        ax.text(x, y, label, ha="center", va="center", color="white", fontsize=8, weight="bold")
    ax.add_patch(plt.Circle((cx, cy), 0.85, color="#0f172a", ec="#22c55e", lw=3))
    ax.text(cx, cy, "Role\nworkspace", ha="center", va="center", color="white", fontsize=10, weight="bold")
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 5.2)
    ax.set_title("Role-based navigation (concept)", fontsize=13, weight="bold")
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="#f1f5f9")
    plt.close(fig)


def _fig_email_workflow(out: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 3.2))
    ax.axis("off")
    steps = [
        (0.8, 1.2, 1.8, 0.9, "Admin sets\nrecipients"),
        (3.0, 1.2, 1.8, 0.9, "Risk score\ncomputed"),
        (5.2, 1.2, 1.8, 0.9, "Daily / emergency\nrules"),
        (7.4, 1.2, 1.8, 0.9, "SMTP\ndelivery"),
    ]
    for x, y, w, h, t in steps:
        ax.add_patch(plt.Rectangle((x, y), w, h, facecolor="#1e293b", ec="#94a3f8", lw=2))
        ax.text(x + w / 2, y + h / 2, t, ha="center", va="center", color="white", fontsize=9, weight="bold")
    for i in range(3):
        ax.arrow(2.65 + i * 2.2, 1.65, 0.28, 0, head_width=0.12, head_length=0.12, fc="#64748b", ec="#64748b")
    ax.set_xlim(0, 10)
    ax.set_ylim(0.8, 2.4)
    ax.set_title("Admin email alerts — workflow (concept)", fontsize=13, weight="bold")
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="#fff7ed")
    plt.close(fig)


def _fig_modules(out: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.axis("off")
    modules = [
        (0.5, 4.5, "Executive\nBriefing"),
        (3.5, 4.5, "Dashboard"),
        (6.5, 4.5, "Global\nSurveillance"),
        (0.5, 2.5, "Uganda\nHotspots"),
        (3.5, 2.5, "Disease\nProfiler"),
        (6.5, 2.5, "Forecast\nLab"),
        (2.0, 0.6, "Action Plan"),
        (5.5, 0.6, "ROI &\nFinancing"),
        (8.2, 0.6, "Admin"),
    ]
    for x, y, t in modules:
        ax.add_patch(plt.Rectangle((x, y), 2.2, 1.1, facecolor="#0c4a6e", ec="#7dd3fc", lw=2))
        ax.text(x + 1.1, y + 0.55, t, ha="center", va="center", color="white", fontsize=9, weight="bold")
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0.3, 6)
    ax.set_title("Main application modules (concept map)", fontsize=13, weight="bold")
    fig.savefig(out, dpi=150, bbox_inches="tight", facecolor="#ecfeff")
    plt.close(fig)


def main() -> None:
    _ensure_deps()
    from docx import Document
    from docx.shared import Inches

    root = Path(__file__).resolve().parents[1]
    out = root / "STI-EPI-FORECAST_Illustrated_Government_Handbook.docx"
    logo = root / "logo1.png"

    with TemporaryDirectory() as tmp:
        td = Path(tmp)
        p1 = td / "fig_arch.png"
        p2 = td / "fig_flow.png"
        p3 = td / "fig_roles.png"
        p4 = td / "fig_email.png"
        p5 = td / "fig_modules.png"
        _fig_architecture(p1)
        _fig_data_flow(p2)
        _fig_roles(p3)
        _fig_email_workflow(p4)
        _fig_modules(p5)

        doc = Document()
        doc.add_heading("STI-EPI-FORECAST — Illustrated government handbook", level=0)
        doc.add_paragraph(
            "This document supplements the comprehensive text handbook with embedded concept diagrams. "
            "Figures are generated for clarity in briefings; they are not live screenshots of the running app."
        )
        if logo.exists():
            doc.add_paragraph("Project branding (logo):")
            doc.add_picture(str(logo), width=Inches(2.4))
            doc.add_paragraph("")

        doc.add_heading("Figure 1 — System architecture", level=1)
        doc.add_picture(str(p1), width=Inches(6.5))
        doc.add_paragraph(
            "The Streamlit dashboard is the primary operator interface. FastAPI provides NLP alerts, "
            "forecasting, and OpenAI-compatible chat proxy endpoints. The Expo app mirrors key surveillance views."
        )

        doc.add_heading("Figure 2 — Open-web signal flow", level=1)
        doc.add_picture(str(p2), width=Inches(6.5))
        doc.add_paragraph(
            "GDELT, Reddit public JSON, Hacker News Algolia, and optional NewsAPI contribute to a cached outbreak snapshot "
            "used across Dashboard, Global Surveillance, and Action Plan social views."
        )

        doc.add_heading("Figure 3 — Role workspace model", level=1)
        doc.add_picture(str(p3), width=Inches(5.8))
        doc.add_paragraph(
            "Six role profiles filter navigation so each user sees modules aligned with incident command, surveillance, "
            "modelling, border operations, policy, or platform administration."
        )

        doc.add_heading("Figure 4 — Admin email alert pipeline", level=1)
        doc.add_picture(str(p4), width=Inches(6.5))
        doc.add_paragraph(
            "Administrators configure recipients and thresholds on the Admin page. The app evaluates risk periodically, "
            "sends scheduled daily summaries, and sends emergency alerts when the risk score exceeds the configured threshold, "
            "subject to cooldown. SMTP credentials are supplied via environment variables."
        )

        doc.add_heading("Figure 5 — Module map", level=1)
        doc.add_picture(str(p5), width=Inches(6.5))
        doc.add_paragraph(
            "Core modules: Executive Briefing, Dashboard, Global Surveillance, Uganda Hotspots, Disease Profiler, "
            "Forecast Lab, Action Plan, ROI & Financing, and Administration."
        )

        doc.add_heading("How to capture real screenshots (optional)", level=1)
        doc.add_paragraph(
            "For a fully photographic manual: run the Streamlit app locally, use Windows Snipping Tool or "
            "Win+Shift+S on each module, save PNG files, and insert them into this document using Word Insert → Pictures. "
            "Recommended captures: sidebar + Dashboard, Global Surveillance tabs, Action Plan tabs, Admin email panel, "
            "FastAPI /docs page, and Expo mobile home screen."
        )

        doc.add_heading("Deployment reminder", level=1)
        doc.add_paragraph(
            "Streamlit Community Cloud: repository MosesLuzinda/STI-EPIFORECAST-DASHBOARD, branch main, main file app.py, "
            "with requirements.txt for dependencies. Configure secrets for AI keys, NewsAPI, and SMTP as needed."
        )

        doc.save(out)

    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
